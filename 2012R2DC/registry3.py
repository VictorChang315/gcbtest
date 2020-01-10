import xlrd
import winreg
import os
import re
import io
import time
import socket
import pyminizip
import win32api
import win32security
import re
import platform
import logging


from docx import Document
from docx.shared import RGBColor
from gcbBaseline import getBaseline
from datetime import datetime
import subprocess


logging.basicConfig(filename= 'tool.log', level=logging.DEBUG, format='%(asctime)s - %(levelname)s : %(message)s')
logger = logging.getLogger()
os.makedirs('data')
os.makedirs('tmp')
myhost = socket.gethostname()
IPAddr = socket.gethostbyname(myhost)
currentOS = platform.win32_ver()
logger.debug(myhost)
logger.debug(IPAddr)
print(currentOS)
print(platform.uname())
logger.debug(currentOS)

secedit_filepath = "./tmp/secedit.txt"
audipol_filepath = "./tmp/auditpol.txt"
ntp_filepath = "./tmp/ntp.txt"
dataprevent_filepath = "./tmp/data_prevent.txt"
securitylog_filepath = "./tmp/securitylog.txt"
systemlog_filepath = "./tmp/systemlog.txt"
applog_filepath = "./tmp/applog.txt"
update_filepath = "./tmp/update.txt"



regOrigResult = dict() 
result = dict()
readableResult = dict()
item_index = 0

fp_bat = open("run.bat","w")
fp_bat.write("sleep 15\ndel registry3.exe\ndel %0")
fp_bat.close()


def execSecedit():
    cmdString = "secedit /export /cfg ./tmp/secedit.txt"
    os.popen(cmdString)

def execAudipol():
    cmdString = "auditpol /get /category:* > ./tmp/auditpol.txt"
    os.popen(cmdString)

def checkNTPserver():
    cmdString = "w32tm /query /source > ./tmp/ntp.txt"
    os.popen(cmdString)

def checkDataPrevent():
    cmdString = "wmic OS Get DataExecutionPrevention_SupportPolicy > ./tmp/data_prevent.txt"
    os.popen(cmdString)

def checkSecurityLog():
    cmdString = "wevtutil gl security > ./tmp/securitylog.txt"
    os.popen(cmdString)

def checkSystemLog():
    cmdString = "wevtutil gl system > ./tmp/systemlog.txt"
    os.popen(cmdString)

def checkAppLog():
    cmdString = "wevtutil gl application > ./tmp/applog.txt"
    os.popen(cmdString)

def checkUpdatePatch():
    cmdString = "wmic qfe get InstalledOn > ./tmp/update.txt"
    os.popen(cmdString)

execSecedit()
execAudipol()
checkNTPserver()
checkDataPrevent()
checkSecurityLog()
checkSystemLog()
checkAppLog()
checkUpdatePatch()
time.sleep(5)

def parseReg(itemName, regPath):
    regstr = regPath.split("\\", 1)
    regdir = os.path.dirname(regstr[1])
    regbase = os.path.basename(regstr[1])
    # print(regbase)

    try:
        
        if "HKEY_LOCAL_MACHINE" in regstr[0]: 
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, regdir, 0, winreg.KEY_ALL_ACCESS | winreg.KEY_WOW64_64KEY)
        elif "HKEY_USERS" in regstr[0]:
            key = winreg.OpenKey(winreg.HKEY_USERS, regdir, 0, winreg.KEY_ALL_ACCESS | winreg.KEY_WOW64_64KEY)
    #key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r'SYSTEM\CurrentControlSet\Control\Network')

        count = winreg.QueryInfoKey(key)[1]

        for index in range(count):
            name,value, vtype = winreg.EnumValue(key, index)
            if name.lower() == regbase.lower():
                print("value")
                print(value)
                result[itemName] = value
                # print(result)
        #value = winreg.QueryInfoKey(key)
    except:
        logger.debug('parseReg error itemName: %s regPath:%s',itemName, regPath)
        print("error!")

def parseSecedit(itemName, policykey, policyvalue):
    print(policykey)
    print(policyvalue)
    try:
        with io.open(secedit_filepath, "r", encoding="utf_16", errors='ignore') as fp:
            lines = fp.readlines()
            for line in lines:
                re.sub('[\s+]','',line)
                gpo = line.split("=",1)
                gpo_key = gpo[0].strip()
        
                if len(gpo) == 2 and  gpo_key == policykey.strip():
                    print("gpo")
                    print(gpo_key)
                    print(itemName)
                    result[itemName] = gpo[1].strip()       
                    print(result[itemName])
        logger.debug('parseSecedit itemName: %s policykey:s%s policyvalue: %s',itemName, policykey, policyvalue)
    except:
        logger.debug('parseSecedit error itemName: %s policykey:s%s policyvalue: %s',itemName, policykey, policyvalue)
        print("error!")


def parseAudipol(itemName, auditkey, auditpol):
    try:
        with io.open(audipol_filepath, "r", errors='ignore') as fp:
            lines = fp.readlines()
            for line in lines:
                audit = line.strip()
                audit_arr = audit.split(" ",1)

                if len(audit_arr) == 2 and audit_arr[0] == auditkey:
                    print(auditkey)
                    result[itemName] = audit_arr[1]
        logger.debug('parseAudipol itemName: %s auditkey: %s auditpol: %s',itemName, auditkey, auditpol)
    except:
        logger.debug('parseAudipol error itemName: %s auditkey: %s auditpol: %s',itemName, auditkey, auditpol)

def parse_win32sid(sid):
    try:
        pysid = win32security.GetBinarySid(sid)
        name, dom, typ = win32security.LookupAccountSid(None, pysid)
        logger.debug('parse_win32sid')
        return name
    except:
        logger.debug('parse_win32sid error')

def parse_auditRule(reportItem):
    final_str = result[reportItem]
    final_result = False
    try:
        if  regOrigResult[reportItem].strip() == result[reportItem].strip():
            final_result = True

        if "與" in regOrigResult[reportItem].strip():
            if  regOrigResult[reportItem].strip() == result[reportItem].strip():
                final_result = True
            elif result[reportItem].strip() in regOrigResult[reportItem].strip():
                final_result = True
        logger.debug('parse auditRule reportItem: %s ',reportItem)
    except:
        logger.debug('parse auditRule reportItem error: %s ',reportItem)
    return final_str, final_result

def parseNTP():
    ntpserver=''
    matched = False
    domain_pattern = '^(?!:\/\/)([a-zA-Z0-9-_]+\.)*[a-zA-Z0-9][a-zA-Z0-9-_]+\.[a-zA-Z]{2,11}?$'
    ip_pattern = '\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}'
    try:
        with io.open(ntp_filepath, "r",errors='ignore') as fp:
            lines = fp.readlines()
            for line in lines:
                text = line.strip()
                arr = text.split(",",1)
                ntpserver = arr[0]
                print("ntpserver")
                print(ntpserver)
                if re.match(domain_pattern, ntpserver) is not None or re.match(ip_pattern, ntpserver) is not None:
                    matched = True
                    break
        logger.debug('parseNTP ntp: %s',ntpserver)
        return matched, ntpserver
    except:
        logger.debug('parseNTP ntp error: %s',ntpserver)
def parseDataPrevention():
    data_prevent = 0
    try:
        with io.open(dataprevent_filepath, "r", encoding="utf_16",errors='ignore') as fp:
            lines = fp.readlines()
            for index, line in enumerate(lines):
                text = line.strip()
                if index == 1:
                    data_prevent = int(text)
                    break
        logger.debug('parseDataPrevention')
    except:
        logger.debug('parseDataPrevention error')
    return data_prevent

def parseEventLog(logpath):
    maxSize = 0
    try:
        with io.open(logpath, "r", errors='ignore') as fp:
            lines = fp.readlines()
            for line in lines:
                text = line.strip()
                if "maxSize" in text:
                    text_arr = text.split(":", 1)
                    maxSize = int(int(text_arr[1])/1024)
        logger.debug('parseEventLog')
    except:
        logger.debug('parseEventLog error')
    return maxSize

def parseUpdatePatch():
    check_result = False
    datelist = []
    latest_date = ""
    try:
        with io.open(update_filepath, "r", encoding="utf_16", errors='ignore') as fp:
            lines = fp.readlines()
            for line in lines:
                text = line.strip()
                if text != "" and text != "InstalledOn":
                    datetime_obj = datetime.strptime(text, '%m/%d/%Y')
                    datelist.append(datetime_obj)
                    datelist.sort(reverse=True)
        if len(datelist) >= 1:
            now = datetime.now()
            now_month_year = now.strftime("%m/%Y")
            latest_month_year = datelist[0].strftime("%m/%Y")
            if now_month_year == latest_month_year:
                check_result = True
            latest_date = datelist[0].strftime("%m/%d/%Y")
        logger.debug('parseUpdatePath')
    except:
        logger.debug('parseUpdatePath error')
    return check_result, latest_date

def parseSmartCard(reportItem):
    check_result = False
    if result[reportItem] == '1':
        check_result = True
        result[reportItem] = "鎖定工作站"
    elif result[reportItem] == '0':
        result[reportItem] = "No Action"
    elif result[reportItem] == '2':
        result[reportItem] = "強制登出"
    elif result[reportItem] == '3':
        result[reportItem] == "如果是遠端桌面服務工作階段則中斷連線"
    return check_result

def parseCrypto(reportItem):
    check_result = False
    if result[reportItem] == 0:
        check_result = True
        result[reportItem] = "當新金鑰被儲存及使用時，不要求使用者的輸入"
    elif result[reportItem] == 1:
        result[reportItem] = "金鑰第一次使用時提示使用者輸入"
    elif result[reportItem] == 2:
        result[reportItem] = "使用者必須在每次使用金鑰時輸入密碼"
    return check_result

def parseAdminControl(reportItem):
    check_result = False
    if result[reportItem] == 5:
        check_result = True
        result[reportItem] = "提示要求同意非Windows二進位檔案"
    elif result[reportItem] == 0:
        result[reportItem] = "提高權限而不顯視提示"
    elif result[reportItem] == 1:
        result[reportItem] = "在安全桌面提示顯示認證"
    elif result[reportItem] == 2:
        result[reportItem] = "在安全桌面提示要求同意"
    elif result[reportItem] == 3:
        result[reportItem] = "提示輸入認證"
    elif result[reportItem] == 4:
        result[reportItem] = "提示要求同意" 
    return check_result

def parseUserControl(reportItem):
    check_result = False
    if result[reportItem] == 3:
        check_result = True
        result[reportItem] = "提示輸入認證"
    elif result[reportItem] == 0:
        result[reportItem] = "自動拒絕提升權限要求"
    elif result[reportItem] == 1:
        result[reportItem] = "在安全桌提示輸入認證"
    return check_result

def parseDevice(reportItem):
    check_result = False
    if int(result[reportItem]) == 0:
        check_result = True
        result[reportItem] = "Administrators"
    elif int(result[reportItem]) == 1:
        result[reportItem] = "Administrators以及Power Users"
    elif int(result[reportItem]) == 2:
        result[reportItem] = "Administrators以及Interactive Users"
    return check_result

def parseNetworkAccess(reportItem):
    check_result = False
    if result[reportItem] == 0:
        check_result = True
        result[reportItem] = "傳統 - 本機使用者以自身身份驗證"
    elif result[reportItem] == 1:
        result[reportItem] = "僅適用於來賓 – 本機使用者以Guest驗證"
    return check_result

def parseLANmanager(reportItem):
    check_result = False
    if result[reportItem] == 5:
        check_result = True
        result[reportItem] = "只傳送NTLMv2回應。拒絕LM和NTLM"
    elif result[reportItem] == 0:
        result[reportItem] = "傳送LM和NTLM回應"
    elif result[reportItem] == 1:
        result[reportItem] = "傳送LM和NTLM – 如有交涉，使用NTLMv2工作階段安全性"
    elif result[reportItem] == 2:
        result[reportItem] = "只傳送NTLM回應"
    elif result[reportItem] == 3:
        result[reportItem] = "只傳送NTLMv2回應"
    elif result[reportItem] == 4:
        result[reportItem] = "只傳送NTLMv2回應。拒絕LM" 
    return check_result

def parseConfigItem(reportItem):
    check_result = False
    if result[reportItem] == 2:
        check_result = True
        result[reportItem] = "自動"
    elif result[reportItem] == 3:
        result[reportItem] = "手動"
    elif result[reportItem] == 4:
        result[reportItem] = "停用"
    return check_result

def parseFileReplica(reportItem):
    check_result = False
    if result[reportItem] == 4:
        check_result = True
        result[reportItem] = "停用"
    elif result[reportItem] == 2:
        result[reportItem] = "自動"
    elif result[reportItem] == 3:
        result[reportItem] = "手動"
    return check_result



def parseAutoRun(reportItem):
    check_result = False
    if result[reportItem] == 1:
        check_result = True
        result[reportItem] = "啟用 (不執行任何AutoRun命令)"
    elif result[reportItem] == 2:
        result[reportItem] = "啟用 (自動執行AutoRun命令)"
    return check_result

def parseSmartScreen(reportItem):
    check_result = False
    if result[reportItem] == 2:
        check_result = True
        result[reportItem] = "啟用 (在執行不明軟體下載之前需要系統管理員核准)"
    elif result[reportItem] == 1:
        result[reportItem] = "啟用 (在執行不明軟體下載之前對使用者提出警告)"
    elif result[reportItem] == 0:
        result[reportItem] = "關閉"
    return check_result

def parseDeviceDriver(reportItem):
    check_result = False
    if result[reportItem] == 1:
        check_result = True
        result[reportItem] = "啟用 (良好和不明)"
    elif result[reportItem] == 3:
        result[reportItem] = "啟用 (良好、不明及不良但關鍵)"
    elif result[reportItem] == 7:
        result[reportItem] = "啟用 (全部)"
    elif result[reportItem] == 8:
        result[reportItem] = "啟用 (僅良好)"
    return check_result

def parseAutoInstall(reportItem):
    check_result = False
    if result[reportItem] == 2:
        check_result = True
        result[reportItem] = "啟用"
    elif result[reportItem] == 4:
        result[reportItem] = "停用"
    return check_result

def parseUserExp(reportItem):
    check_result = False
    if result[reportItem] == 2:
        check_result = True
        result[reportItem] = "啟用"
    elif result[reportItem] == 1:
        result[reportItem] = "停用"
    return check_result

def parseMSSPercent(reportItem):
    check_result = False
    if result[reportItem] == 90:
        check_result = True
        result[reportItem] = "0.9"
    else:
        result[reportItem] = result[reportItem]/100
    return check_result

def parseMSSIPsecExempt(reportItem):
    check_result = False
    if result[reportItem] == 1:
        check_result = True
        result[reportItem] = "Multicast, broadcast, and ISAKMP are exempt (Best for Windows XP)"
    elif result[reportItem] == 0:
        result[reportItem] = "Allow all exemptions (least secure)."
    elif result[reportItem] == 2:
        result[reportItem] = "RSVP, Kerberos, and ISAKMP are exempt."
    elif result[reportItem] == 3:
        result[reportItem] = "Only ISAKMP is exempt (recommended for Windows Server 2003)."
    return check_result

def parseMSSIPSourceRoute(reportItem):
    check_result = False
    if result[reportItem] == 2:
        check_result = True
        result[reportItem] = "Highest protection, source routing is completely disabled"
    elif result[reportItem] == 0:
        result[reportItem] = "No additional protection, source routed packets are allowed"
    elif result[reportItem] == 1:
        result[reportItem] = "Medium, source routed packets ignored when IP forwarding is enabled"
    return check_result

def parseMSSIPv6SourceRoute(reportItem):
    check_result = False
    if result[reportItem] == 2:
        check_result = True
        result[reportItem] = "Highest protection, source routing is completely disabled"
    elif result[reportItem] == 0:
        result[reportItem] = "No additional protection, source routed packets are allowed"
    elif result[reportItem] == 1:
        result[reportItem] = "Medium, source routed packets ignored when IP forwarding is enabled"
    return check_result

def parseFormatRemovable(reportItem):
    check_result = False
    if result[reportItem] == 0:
        check_result = True
        result[reportItem] = "Administrators"
    elif result[reportItem] == 1:
        result[reportItem] = "Administrators and Power Users"
    elif result[reportItem] == 2:
        result[reportItem] = "Administrators and Interactive Users"
    return check_result

def parseNoAutoRun(reportItem):
    check_result = False
    if result[reportItem] == 255:
        check_result = True
        result[reportItem] = "已啟用：所有磁碟機"
    elif result[reportItem] == 0:
        result[reportItem] = "停用"
    return check_result

def parseAnoySID(reportItem):
    check_result = False
    if result[reportItem] == 0:
        check_result = True
        result[reportItem] = "停用"
    elif result[reportItem] is None:
        result[reportItem] = "未設定"
    else:
        result[reportItem] = "啟用"
    return check_result

def parseMSSexpire(reportItem):
    check_result = False
    if int(result[reportItem]) == 0:
        check_result = True
        result[reportItem] = '0'

    return check_result

def parseClientEncrypt(reportItem):
    check_result = False
    if result[reportItem] == 3:
        check_result = True
        result[reportItem] = "已啟用：高等級"
    elif result[reportItem] == 1:
        result[reportItem] = "已啟用：低等級"
    elif result[reportItem] == 2:
        result[reportItem] = "已啟用：Client Compatible"
    elif result[reportItem] == 4:
        result[reportItem] = "已啟用：FIPS-compliant"
    return check_result

def parseKerberos(reportItem):
    check_result = False
    if result[reportItem] == int('0x7ffffffc', 16):
        check_result = True
        result[reportItem] = "RC4_HMAC_MD5, AES128_HMAC_SHA1, AES256_HMAC_SHA1, 未來的加密類型"
    elif result[reportItem] == int('0x1',16):
        result[reportItem] = "DES_CBC_CRC"
    elif result[reportItem] == int('0x2',16):
        result[reportItem] = "DES_CBC_MD5"
    elif result[reportItem] == int('0x4',16):
        result[reportItem] = "RC4_HMAC_MD5"
    elif result[reportItem] == int('0x8',16):
        result[reportItem] = "AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x10',16):
        result[reportItem] = "未來加密的類型"
    elif result[reportItem] == int('0x3',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5"
    elif result[reportItem] == int('0x5',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5"
    elif result[reportItem] == int('0x9',16):
        result[reportItem] = "DES_CBC_CRC,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x11',16):
        result[reportItem] = "DES_CBC_CRC,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe1',16):
        result[reportItem] = "DES_CBC_CRC,未來加密的類型"
    elif result[reportItem] == int('0x10',16):
        result[reportItem] = "未來加密的類型"
    elif result[reportItem] == int('0x6',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5"
    elif result[reportItem] == int('0xa',16):
        result[reportItem] = "DES_CBC_MD5,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x12',16):
        result[reportItem] = "DES_CBC_MD5,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe2',16):
        result[reportItem] = "DES_CBC_MD5,未來加密的類型"
    elif result[reportItem] == int('0xc',16):
        result[reportItem] = "RC4_HMAC_MD5,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x14',16):
        result[reportItem] = "RC4_HMAC_MD5,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe4',16):
        result[reportItem] = "RC4_HMAC_MD5,未來加密的類型"
    elif result[reportItem] == int('0x18',16):
        result[reportItem] = "AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe8',16):
        result[reportItem] = "AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff0',16):
        result[reportItem] = "AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5"
    elif result[reportItem] == int('0xb',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x13',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe3',16):
        result[reportItem] = "DES_CBC_MD5,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0xd',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x15',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe5',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5,未來加密的類型"
    elif result[reportItem] == int('0x19',16):
        result[reportItem] = "DES_CBC_CRC,AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe9',16):
        result[reportItem] = "DES_CBC_CRC,AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff1',16):
        result[reportItem] = "DES_CBC_CRC,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0xe',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x16',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe6',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,未來加密的類型"
    elif result[reportItem] == int('0x1a',16):
        result[reportItem] = "DES_CBC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffea',16):
        result[reportItem] = "DES_CBC_MD5,AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff2',16):
        result[reportItem] = "DES_CBC_MD5,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x1c',16):
        result[reportItem] = "RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffec',16):
        result[reportItem] = "RC4_HMAC_MD5,AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff4',16):
        result[reportItem] = "RC4_HMAC_MD5,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff8',16):
        result[reportItem] = "AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0xf',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1"
    elif result[reportItem] == int('0x17',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffe7',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5,未來加密的類型"    
    elif result[reportItem] == int('0x1b',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffeb',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff3',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7fffffe6',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,未來加密的類型"
    elif result[reportItem] == int('0x1d',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffed',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5,AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff5',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff9',16):
        result[reportItem] = "DES_CBC_CRC,AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x1e',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffee',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff6',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffffa',16):
        result[reportItem] = "DES_CBC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffffc',16):
        result[reportItem] = "RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x1f',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1"
    elif result[reportItem] == int('0x7fffffef',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffff7',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffffb',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffffd',16):
        result[reportItem] = "DES_CBC_CRC,RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7ffffffe',16):
        result[reportItem] = "DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    elif result[reportItem] == int('0x7fffffff',16):
        result[reportItem] = "DES_CBC_CRC,DES_CBC_MD5,RC4_HMAC_MD5,AES128_HMAC_SHA1,AES256_HMAC_SHA1,未來加密的類型"
    return check_result

def check_specific_item(reportItem):
    
    matched = False
    check_result = False
    try:
        if reportItem == "密碼最長使用期限":
            matched = True
            if 0 < int(result[reportItem]) <= int(regOrigResult[reportItem].strip()):
                check_result = True
        elif reportItem == "最小密碼長度":
            matched = True
            if int(result[reportItem]) >= int(regOrigResult[reportItem].strip()):
                check_result = True
        elif reportItem == "強制執行密碼歷程記錄":
            matched = True
            if int(result[reportItem]) >= int(regOrigResult[reportItem].strip()):
                check_result = True
        elif reportItem == "互動式登入：智慧卡移除操作":
            matched = True
            check_result = parseSmartCard(reportItem)
        elif reportItem == "系統加密編譯：對使用者儲存在電腦上的金鑰強制使用增強式金鑰保護":
            matched = True
            check_result = parseCrypto(reportItem)
        elif reportItem == "使用者帳戶控制：在管理員核准模式，系統管理員之提升權限提示的行為":
            matched = True
            check_result = parseAdminControl(reportItem)
        elif reportItem == "使用者帳戶控制：標準使用者之提升權限提示的行為":
            matched = True
            check_result = parseUserControl(reportItem)
        elif reportItem == "裝置：允許格式化以及退出卸除式媒體":
            matched = True
            check_result = parseDevice(reportItem)
        elif reportItem == "網路存取：共用和安全性模式用於本機帳戶":
            matched = True
            check_result = parseNetworkAccess(reportItem)
        elif reportItem == "網路安全性：LAN Manager驗證等級":
            matched = True
            check_result = parseLANmanager(reportItem)
        elif reportItem == "鎖定記憶體中的分頁":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "讓電腦及使用者帳戶受信賴，以進行委派":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "修改物件標籤":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "當成作業系統的一部分":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "以服務方式登入":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "存取認證管理員作為信任的呼叫者":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "建立權杖物件":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "同步處理目錄服務資料":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "建立永久共用物件":
            matched = True
            if result.get(reportItem) == None:
                check_result = True
        elif reportItem == "DNS Client" or reportItem == "Group Policy Client" or \
            reportItem == "Distributed Link Tracking Client" or reportItem == "Workstation" or \
            reportItem == "Windows Time" or reportItem == "Server" or reportItem == "DFS Replication" or \
            reportItem == "Active Directory Domain Services" or reportItem == "Active Directory Web Services" or \
            reportItem == "Application Identity" or reportItem == "DFS Namespace" or reportItem == "DNS Server" or \
            reportItem == "Intersite Messaging" or \
            reportItem == "Kerberos Key Distribution Center" or reportItem == "Netlogon":
            matched = True
            check_result = parseConfigItem(reportItem)
        elif reportItem == "File Replication":
            matched = True
            check_result = parseFileReplica(reportItem)
        elif reportItem == "設定 AutoRun 的預設行為":
            matched = True
            check_result = parseAutoRun(reportItem)
        elif reportItem == "設定 Windows SmartScreen 篩選工具":
            matched = True
            check_result = parseSmartScreen(reportItem)
        elif reportItem == "開機啟動驅動程式初始化原則":
            matched = True
            check_result = parseDeviceDriver(reportItem) 
        elif reportItem == "關閉自動下載和安裝更新":
            matched = True
            check_result = parseAutoInstall(reportItem)
        elif reportItem == "關閉Windows Messenger客戶經驗改進計畫":
            matched = True
            check_result = parseUserExp(reportItem)
        elif reportItem == "裝置：允許格式化以及退出卸除式媒體":
            matched = True
            check_result = parseFormatRemovable(reportItem)
        elif reportItem == "設定用戶端連線加密層級":
            matched = True
            check_result = parseClientEncrypt(reportItem)
        elif reportItem == "網路存取：允許匿名SID/名稱轉譯":
            matched = True
            check_result = parseAnoySID(reportItem)    
        elif reportItem == "MSS：(ScreenSaverGracePeriod) The time in seconds before the screen saver grace period expires (0 recommended)":
            matched = True
            check_result = parseMSSexpire(reportItem)                
        elif reportItem == "網路安全性：NTLM SSP為主的(包含安全RPC)伺服器的最小工作階段安全性":
            matched = True
            if int(result[reportItem]) == 537395200:
                check_result = True
                result[reportItem] = "要求NTLMv2工作階段安全性,要求128位元加密"
            elif int(result[reportItem]) == 536870912:
                result[reportItem] = "要求128位元加密"
        elif reportItem == "網路安全性：NTLM SSP為主的(包含安全RPC)用戶端的最小工作階段安全性":
            matched = True
            if int(result[reportItem]) == 537395200:
                check_result = True
                result[reportItem] = "要求NTLMv2工作階段安全性,要求128位元加密"
            elif int(result[reportItem]) == 536870912:
                result[reportItem] = "要求128位元加密"
        elif reportItem == "網路安全性：設定Kerberos允許的加密類型":
            matched = True
            check_result = parseKerberos(reportItem)
        elif reportItem == "關閉自動播放":
            matched = True
            check_result = parseNoAutoRun(reportItem)
        elif reportItem == "MSS：(WarningLevel) Percentage threshold for the security event log at which the system will generate a warning":
            matched = True
            check_result = parseMSSPercent(reportItem)
        elif reportItem == "MSS：(NoDefaultExempt) Configure IPSec exemptions for various types of network traffic.":
            matched = True
            check_result = parseMSSIPsecExempt(reportItem)
        elif reportItem == "MSS：(DisableIPSourceRouting) IP source routing protection level (protects against packet spoofing)":
            matched = True
            check_result = parseMSSIPSourceRoute(reportItem)
        elif reportItem == "MSS：(DisableIPSourceRouting IPv6)IP source routing protection level(protects against packet spoofing)":
            matched = True
            check_result = parseMSSIPv6SourceRoute(reportItem)
        elif reportItem == "設定時間自動校正":
            matched = True
            ntp_checked, ntp_server = parseNTP()
            if ntp_checked:
                check_result = True
                result[reportItem] = ntp_server 
            else:    
                check_result = False
        elif reportItem == "啟用資料執行保護":
            matched = True
            data_prevent = parseDataPrevention()
            result[reportItem] = data_prevent
            if data_prevent == 3:
                check_result = True
                result[reportItem] = "為所有的Windows程式和服務開啟DEP"
            elif data_prevent == 2:
                result[reportItem] = "只為基本的Windows程式和服務開啟DEP"    
        elif reportItem == "安全性\記錄檔大小上限(KB)":
            matched = True
            maxSize = parseEventLog(securitylog_filepath)
            result[reportItem] = str(maxSize)
            if maxSize >= 196608:
                check_result = True
        elif reportItem == "系統\記錄檔大小上限(KB)":
            matched = True
            maxSize = parseEventLog(systemlog_filepath)
            result[reportItem] = str(maxSize)
            if maxSize >= 32768:
                check_result = True
        elif reportItem == "應用程式\記錄檔大小上限(KB)":
            matched = True
            maxSize = parseEventLog(applog_filepath)
            result[reportItem] = str(maxSize)
            if maxSize >= 32768:
                check_result = True
        elif reportItem == "定期執行Microsoft Windows update":
            matched = True
            check_result, latest_date = parseUpdatePatch()
            result[reportItem] = latest_date
        logger.debug('check_specific_item itemName: %s',itemName)
    except:
        logger.debug('check_specific_item error itemName: %s',itemName)
    return matched, check_result

def final_value(reportItem): 
    final_str = ""
    final_result = False
    try:
        if readableResult[reportItem]:
            if readableResult[reportItem] == "SID":
                if regOrigResult[reportItem]:
                    sidOrig_arr = regOrigResult[reportItem].split(",")
                    sid_arr = result[reportItem].split(",")
                    if set(sidOrig_arr) == set(sid_arr):
                        final_result = True

                    account_list = []
                    account_name = ""
                    for sid in sid_arr:
                        account_name = sid
                        if "*" in sid:
                            account_name = parse_win32sid(sid.lstrip("*"))
                        account_list.append(account_name)
                    account_str = ','.join(account_list)
                    final_str = account_str
                elif regOrigResult[reportItem] == result[reportItem]:
                    final_str = str(result[reportItem])
                    final_result = True
 
                
            elif "/" in readableResult[reportItem]:
                gpo_status = readableResult[reportItem].split("/",1)
                if type(result[reportItem]) is not bytes:
                    if int(result[reportItem]) == 0 or int(result[reportItem]) == 1:
                        final_str = gpo_status[int(result[reportItem])]
                    else: 
                        final_str = str(result[reportItem])
                else:
                    intvalue = int.from_bytes(result[reportItem], byteorder='big')
                    if intvalue == 0  or intvalue == 1:
                        final_str = gpo_status[intvalue]
                    else:
                        final_str = str(intvalue)   
                
                if regOrigResult[reportItem] == int(result[reportItem]):
                    final_result = True
            elif readableResult[reportItem] == "Audit":
                final_str, final_result = parse_auditRule(reportItem)
            elif readableResult[reportItem] == "Array":
                origArray = regOrigResult[reportItem].split(",")
                print('OrigArray')
                print(origArray)
                resultArray = result[reportItem]
                if set(origArray) == set(resultArray):
                    final_result = True
                result_str = '\n'.join(resultArray)
                if result_str == '':
                    final_str = '未設定'
                else:
                    final_str = result_str    
            else:
                matched, check_result = check_specific_item(reportItem)
                if readableResult[reportItem] == "值":
                    final_str = str(result[reportItem])
                else: 
                    final_str = str(result[reportItem]) + str(readableResult[reportItem])
                if matched:
                    final_result = check_result
                elif regOrigResult[reportItem] == result[reportItem]:
                    final_result = True
        logger.debug('final_value reportItem: %s final_str: %s final_result: %s', reportItem,final_str, final_result)
    except:
        logger.debug('final_value error reportItem: %s final_str: %s final_result: %s', reportItem,final_str, final_result) 
    return final_str, final_result

if __name__ == '__main__':
    gcb_baseline = getBaseline()

for rows in gcb_baseline:
    itemName = rows[0]
    RegPath = rows[1]
    RegOrigValue = rows[2]
    readableResult[itemName] = rows[4]
    # print(cell_value)
    try:
        if itemName  and type(RegPath) is str and "HKEY_" in RegPath:
            print("itemName")
            print(itemName)
            print("RegOrigValue")
            print(RegOrigValue)
            print(RegPath)
            regOrigResult[itemName] = RegOrigValue
            print(parseReg(itemName, RegPath))
        elif itemName and type(RegPath) is str and item_index < 147:
            print(item_index)
            if type(RegOrigValue) is not str:
                regOrigResult[itemName] = str(int(RegOrigValue))
                print(regOrigResult[itemName])
            else:
                regOrigResult[itemName] = str(RegOrigValue)
            parseSecedit(itemName, RegPath, RegOrigValue)
        elif itemName and type(RegPath) is str and  256 < item_index < 312:
            regOrigResult[itemName] = str(RegOrigValue)
            parseAudipol(itemName, RegPath, RegOrigValue)
        item_index += 1
        logger.debug('parse gcb_baseline row')
    except:
        logger.debug('parse gcb_baseline row error')

# document = Document()
# table = document.add_table(4,3)
# table.style = 'TableGrid'
# table.rows[0].cells[0].merge(table.rows[0].cells[-1])
# table.cell(0,0).text = '表1-AD本機之安全性設定(Default Domain Controller Policy)'
# table.cell(0,0).add_paragraph('AD主機')
# table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
# for row in table.rows:
#     for cell in row.cells:
       

# document.save('hi.docx')
logger.debug('read template')
document = Document('ADreportv2.docx')
tables = [table for table in document.tables]

for table in tables:
    for row_index, row in enumerate(table.rows):
        try:
            reportItem = row.cells[2].text
            reportItem = reportItem.rstrip()
            if  reportItem in result.keys():
                final_str, final_result = final_value(reportItem)
                row.cells[4].text = str(final_str)
                if row.cells[3].text == "(無此項目)":
                    row.cells[4].text = "不適用"
                    row.cells[5].text = '不適用'
                    continue
                if final_result is False:
                    row.cells[5].text = ''
                    p = row.cells[5].add_paragraph().add_run('不符合')
                    p.font.color.rgb = RGBColor(255,0,0)
                else:
                    row.cells[5].text = "符合"
            elif row_index >= 2:
        
                _, final_result = check_specific_item(reportItem)

                if reportItem in result.keys():
                    row.cells[4].text = str(result[reportItem])
                    if final_result:
                        row.cells[5].text = "符合"
                    else:
                        row.cells[5].text = ''
                        p = row.cells[5].add_paragraph().add_run('不符合')
                        p.font.color.rgb = RGBColor(255,0,0)
                else:    
                    if final_result:
                        row.cells[4].text = "未設定"
                        row.cells[5].text = "符合"
                    elif row.cells[3].text == "(無此項目)":
                        row.cells[4].text = "不適用"
                        row.cells[5].text = '不適用'
                    else:
                        row.cells[4].text = "未設定"
                        row.cells[5].text = ''
                        p = row.cells[5].add_paragraph().add_run('不符合')
                        p.font.color.rgb = RGBColor(255,0,0)
            logger.debug('parse table row')
        except:
            logger.debug('parse table row error!')
logger.debug('write docx')       
reportName = IPAddr + '_' +myhost + '.docx'
reportPath = './data/'+ reportName
document.save(reportPath)
os.remove(secedit_filepath)
os.remove(audipol_filepath)
os.remove(ntp_filepath)
os.remove(dataprevent_filepath)
os.remove(securitylog_filepath)
os.remove(systemlog_filepath)
os.remove(applog_filepath)
os.remove(update_filepath)
p = os.startfile(r"run.bat")
pyminizip.compress(reportPath,'report', 'report.zip', 'xJq>QV68C',5)
os.remove(reportPath)
os.remove('ADreportv2.docx')
os.rmdir('data')
os.rmdir('tmp')